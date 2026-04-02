from pathlib import Path
import csv

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT_CSV = ROOT / "artifacts" / "statistics_336" / "table_5_1_baseline_test_results.csv"
OUTPUT_DOCX = ROOT / "artifacts" / "statistics_336" / "table_5_1_baseline_test_results_word.docx"
OUTPUT_TXT = ROOT / "artifacts" / "statistics_336" / "table_5_1_baseline_test_results_word_ready.txt"

CAPTION = (
    "Table 5.1. Baseline model performance on the held-out test set of Dataset #336. "
    "Values are reported as point estimates with 95% bootstrap confidence intervals."
)
NOTE = (
    "Note. Values are presented as point estimate (95% bootstrap CI) based on 2,000 "
    "stratified percentile bootstrap replicates. HistGradientBoostingClassifier was used "
    "as the boosting-oriented baseline in place of XGBoost because of an environment-"
    "specific compatibility issue with the local software stack."
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    font = run.font
    font.name = "Times New Roman"
    font.size = Pt(size)
    # Needed for Word to respect East Asia font mapping.
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def load_rows() -> tuple[list[str], list[dict[str, str]]]:
    with INPUT_CSV.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        return reader.fieldnames or [], list(reader)


def export_txt(headers: list[str], rows: list[dict[str, str]]) -> None:
    lines = [CAPTION, "", "\t".join(headers)]
    for row in rows:
        lines.append("\t".join(row[h] for h in headers))
    lines.extend(["", NOTE])
    OUTPUT_TXT.write_text("\n".join(lines), encoding="utf-8")


def build_docx(headers: list[str], rows: list[dict[str, str]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_run = caption.add_run(CAPTION)
    caption_run.bold = True
    caption_run.font.name = "Times New Roman"
    caption_run.font.size = Pt(12)
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, size=10)
        set_cell_shading(cell, "D9E2F3")

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, header in enumerate(headers):
            set_cell_text(table.rows[row_idx].cells[col_idx], row[header], size=10)

    note = doc.add_paragraph()
    note_format = note.paragraph_format
    note_format.space_before = Pt(8)
    note_run = note.add_run(NOTE)
    note_run.italic = True
    note_run.font.name = "Times New Roman"
    note_run.font.size = Pt(10)
    note_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.save(OUTPUT_DOCX)


def main() -> None:
    headers, rows = load_rows()
    export_txt(headers, rows)
    build_docx(headers, rows)
    print(OUTPUT_DOCX)
    print(OUTPUT_TXT)


if __name__ == "__main__":
    main()
